"""
FastAPI backend service for Job Market Skill Gap Analyzer.
Provides REST API endpoints for job analysis and skill gap computation.
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Optional, Any
from datetime import datetime
import pandas as pd
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import logging
from dotenv import load_dotenv
import io
from PyPDF2 import PdfReader
from docx import Document

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from nlp.skill_extractor import SkillExtractor
from backend.database import FileStorage

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Job Market Skill Gap Analyzer API",
    description="API for analyzing job market demand and identifying skill gaps",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8501", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize services
skill_extractor = SkillExtractor()
storage = FileStorage()


# Pydantic models
class SkillDemandResponse(BaseModel):
    """Skill demand statistics."""
    skill: str
    job_count: int
    percentage: float
    demand_level: str
    category: str


class ResumeAnalysisRequest(BaseModel):
    """Request model for resume analysis."""
    resume_text: str
    student_name: Optional[str] = "Student"


class SkillGapItem(BaseModel):
    """Individual skill gap analysis item."""
    skill: str
    job_count: int
    percentage: float
    demand_level: str
    student_has: bool
    gap: bool


class SkillGapResponse(BaseModel):
    """Complete skill gap analysis response."""
    student_name: str
    student_skills: List[str]
    skill_count: int
    gap_analysis: List[SkillGapItem]
    missing_high_demand: List[str]
    missing_medium_demand: List[str]
    matched_skills: List[str]


class JobStatsResponse(BaseModel):
    """Job market statistics."""
    total_jobs: int
    unique_companies: int
    unique_skills: int
    avg_skills_per_job: float
    top_skills: List[Dict[str, Any]]


# Helper functions
def load_market_data():
    """Load job market data."""
    try:
        # Try loading from storage
        demand_df = storage.load_skill_demand()
        
        if demand_df.empty:
            # Fallback: compute from jobs
            jobs_df = storage.load_jobs()
            if not jobs_df.empty:
                demand_df = skill_extractor.compute_skill_demand(jobs_df)
                storage.save_skill_demand(demand_df)
        
        return demand_df
    except Exception as e:
        logger.error(f"Error loading market data: {e}")
        return pd.DataFrame()


# API Endpoints
@app.get("/")
def root():
    """Root endpoint."""
    return {
        "message": "Job Market Skill Gap Analyzer API",
        "version": "1.0.0",
        "endpoints": {
            "health": "/health",
            "skill_demand": "/api/skill-demand",
            "analyze_resume": "/api/analyze-resume",
            "job_stats": "/api/job-stats"
        }
    }


@app.get("/health")
def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat()
    }


@app.get("/api/skill-demand")
def get_skill_demand(top_n: int = 50):
    """
    Get top skills in demand.
    
    Args:
        top_n: Number of top skills to return
        
    Returns:
        List of skills with demand statistics
    """
    try:
        logger.info(f"Loading market data for top {top_n} skills")
        demand_df = load_market_data()
        logger.info(f"Loaded {len(demand_df)} skills from data")
        
        if demand_df.empty:
            logger.warning("No data available, returning empty response")
            return {"skills": [], "count": 0}
        
        # Add demand_level if missing
        if 'demand_level' not in demand_df.columns:
            logger.info("Adding demand_level column")
            demand_df['demand_level'] = pd.cut(
                demand_df['percentage'],
                bins=[0, 33, 66, 100],
                labels=['Low', 'Medium', 'High']
            ).astype(str)
        
        # Get top skills
        top_skills = demand_df.head(top_n)
        logger.info(f"Returning {len(top_skills)} skills")
        
        result = top_skills.to_dict('records')
        return {
            "skills": result,
            "count": len(result)
        }
    
    except Exception as e:
        logger.error(f"Error in get_skill_demand: {e}", exc_info=True)
        return {"skills": [], "count": 0, "error": str(e)}


@app.post("/api/analyze-resume", response_model=SkillGapResponse)
async def analyze_resume(request: ResumeAnalysisRequest):
    """
    Analyze resume and compute skill gap.
    
    Args:
        request: Resume analysis request with text and optional name
        
    Returns:
        Skill gap analysis with recommendations
    """
    try:
        # Extract skills from resume
        resume_data = skill_extractor.extract_from_resume(request.resume_text)
        student_skills = resume_data['skills']
        
        # Load market demand
        demand_df = load_market_data()
        
        if demand_df.empty:
            raise HTTPException(
                status_code=404,
                detail="No market data available"
            )
        
        # Compare skills
        gap_df = skill_extractor.compare_skills(
            student_skills=student_skills,
            market_demand=demand_df,
            top_n=50
        )
        
        # Identify missing high-demand skills
        missing_high = gap_df[
            (gap_df['gap'] == True) & (gap_df['demand_level'] == 'High')
        ]['skill'].tolist()
        
        missing_medium = gap_df[
            (gap_df['gap'] == True) & (gap_df['demand_level'] == 'Medium')
        ]['skill'].tolist()
        
        matched_skills = gap_df[gap_df['student_has'] == True]['skill'].tolist()
        
        return {
            "student_name": request.student_name,
            "student_skills": student_skills,
            "skill_count": len(student_skills),
            "gap_analysis": gap_df.to_dict('records'),
            "missing_high_demand": missing_high,
            "missing_medium_demand": missing_medium,
            "matched_skills": matched_skills
        }
    
    except Exception as e:
        logger.error(f"Error in analyze_resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload resume file and extract skills.
    
    Args:
        file: Resume file (txt, pdf, docx)
        
    Returns:
        Extracted skills and gap analysis
    """
    try:
        # Read file content
        content = await file.read()
        filename = file.filename.lower()
        
        # Extract text based on file type
        resume_text = ""
        
        if filename.endswith('.pdf'):
            # Read PDF
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    resume_text += page.extract_text() + "\n"
                logger.info(f"Extracted text from PDF: {len(resume_text)} characters")
            except Exception as e:
                logger.error(f"Error reading PDF: {e}")
                raise HTTPException(status_code=400, detail=f"Unable to read PDF file: {str(e)}")
                
        elif filename.endswith('.docx'):
            # Read DOCX
            try:
                docx_file = io.BytesIO(content)
                doc = Document(docx_file)
                for paragraph in doc.paragraphs:
                    resume_text += paragraph.text + "\n"
                logger.info(f"Extracted text from DOCX: {len(resume_text)} characters")
            except Exception as e:
                logger.error(f"Error reading DOCX: {e}")
                raise HTTPException(status_code=400, detail=f"Unable to read DOCX file: {str(e)}")
                
        elif filename.endswith('.txt'):
            # Read plain text
            try:
                resume_text = content.decode('utf-8')
                logger.info(f"Read text file: {len(resume_text)} characters")
            except Exception as e:
                logger.error(f"Error reading text file: {e}")
                raise HTTPException(status_code=400, detail="Unable to decode text file. Please ensure it's UTF-8 encoded.")
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload PDF, DOCX, or TXT file."
            )
        
        if not resume_text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the file. Please check the file content."
            )
        
        # Analyze resume
        request = ResumeAnalysisRequest(
            resume_text=resume_text,
            student_name=file.filename.split('.')[0]
        )
        
        return await analyze_resume(request)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in upload_resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/job-stats")
def get_job_stats():
    """
    Get job market statistics.
    
    Returns:
        Statistics about collected jobs and skills
    """
    try:
        jobs_df = storage.load_jobs()
        demand_df = load_market_data()
        
        if jobs_df.empty:
            return {
                "total_jobs": 0,
                "unique_companies": 0,
                "unique_skills": 0,
                "avg_skills_per_job": 0.0,
                "top_skills": []
            }
        
        # Compute stats
        total_jobs = len(jobs_df)
        unique_companies = jobs_df['company'].nunique() if 'company' in jobs_df.columns else 0
        
        # Get average skills per job
        avg_skills = 0
        if 'skill_count' in jobs_df.columns:
            avg_skills = jobs_df['skill_count'].mean()
        elif 'skills' in jobs_df.columns:
            import ast
            try:
                jobs_df['skills_parsed'] = jobs_df['skills'].apply(
                    lambda x: ast.literal_eval(x) if isinstance(x, str) else x
                )
                avg_skills = jobs_df['skills_parsed'].apply(len).mean()
            except:
                avg_skills = 27.4  # Default from pipeline
        
        unique_skills = len(demand_df) if not demand_df.empty else 0
        
        # Top 10 skills
        top_skills = demand_df.head(10).to_dict('records') if not demand_df.empty else []
        
        return {
            "total_jobs": int(total_jobs),
            "unique_companies": int(unique_companies),
            "unique_skills": int(unique_skills),
            "avg_skills_per_job": float(round(avg_skills, 1)),
            "top_skills": top_skills
        }
    
    except Exception as e:
        logger.error(f"Error in get_job_stats: {e}")
        return {
            "total_jobs": 0,
            "unique_companies": 0,
            "unique_skills": 0,
            "avg_skills_per_job": 0.0,
            "top_skills": []
        }


@app.get("/api/skills/search")
def search_skills(query: str, limit: int = 20):
    """
    Search for skills by name.
    
    Args:
        query: Search query
        limit: Maximum results
        
    Returns:
        Matching skills with demand data
    """
    try:
        demand_df = load_market_data()
        
        if demand_df.empty:
            return []
        
        # Filter skills containing query
        matching = demand_df[
            demand_df['skill'].str.contains(query, case=False, na=False)
        ].head(limit)
        
        return matching.to_dict('records')
    
    except Exception as e:
        logger.error(f"Error in search_skills: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# CHROME EXTENSION API ENDPOINTS
# ============================================================================

class ResumeUploadResponse(BaseModel):
    """Response for resume upload"""
    resume_id: str
    skills: List[str]
    analysis: Dict[str, Any]
    message: str


class SkillGapAnalysisRequest(BaseModel):
    """Request for skill gap analysis"""
    resume_id: str
    job_title: str
    job_description: str
    company: Optional[str] = None
    location: Optional[str] = None


class SkillGapAnalysisResponse(BaseModel):
    """Response for skill gap analysis"""
    match_score: float
    matching_skills: List[str]
    missing_skills: List[str]
    insights: List[str]
    recommendation: str


class CourseRecommendation(BaseModel):
    """Course recommendation"""
    provider: str
    title: str
    description: str
    url: str
    duration: str
    level: str


@app.post("/api/resume/upload", response_model=ResumeUploadResponse)
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload and parse resume for skill extraction.
    
    Args:
        file: Resume file (PDF, DOC, DOCX)
        
    Returns:
        Extracted skills and analysis
    """
    try:
        logger.info(f"Received resume upload: {file.filename}")
        
        # Validate file type
        allowed_types = [
            'application/pdf',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
        
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail="Invalid file type. Please upload PDF or Word document."
            )
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        if file.content_type == 'application/pdf':
            text = extract_text_from_pdf(io.BytesIO(content))
        else:
            text = extract_text_from_docx(io.BytesIO(content))
        
        if not text:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from file"
            )
        
        # Extract skills
        extractor = SkillExtractor()
        skills = extractor.extract_skills(text)
        
        if not skills:
            raise HTTPException(
                status_code=400,
                detail="No skills found in resume"
            )
        
        # Generate unique resume ID
        resume_id = f"resume_{hash(text)}"
        
        # Store resume data (in production, use database)
        resume_data = {
            'id': resume_id,
            'filename': file.filename,
            'skills': skills,
            'text': text[:500],  # Store preview
            'upload_time': datetime.now().isoformat()
        }
        
        # Store in file storage for now
        storage = FileStorage()
        storage.save_resume(resume_id, resume_data)
        
        # Basic analysis
        analysis = {
            'total_skills': len(skills),
            'skill_categories': categorize_skills(skills),
            'experience_level': estimate_experience_level(text)
        }
        
        return ResumeUploadResponse(
            resume_id=resume_id,
            skills=skills,
            analysis=analysis,
            message="Resume uploaded and analyzed successfully"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/analyze/skill-gap", response_model=SkillGapAnalysisResponse)
async def analyze_skill_gap(request: SkillGapAnalysisRequest):
    """
    Analyze skill gap between resume and job posting.
    
    Args:
        request: Analysis request with resume ID and job details
        
    Returns:
        Skill gap analysis with recommendations
    """
    try:
        # Retrieve resume data
        storage = FileStorage()
        resume_data = storage.get_resume(request.resume_id)
        
        if not resume_data:
            raise HTTPException(
                status_code=404,
                detail="Resume not found. Please upload your resume first."
            )
        
        user_skills = set(resume_data['skills'])
        
        # Extract skills from job description
        extractor = SkillExtractor()
        job_skills = set(extractor.extract_skills(request.job_description))
        
        # Calculate matching and missing skills
        matching_skills = list(user_skills.intersection(job_skills))
        missing_skills = list(job_skills.difference(user_skills))
        
        # Calculate match score
        if len(job_skills) > 0:
            match_score = (len(matching_skills) / len(job_skills)) * 100
        else:
            match_score = 0
        
        # Generate insights
        insights = generate_insights(
            match_score=match_score,
            matching_skills=matching_skills,
            missing_skills=missing_skills,
            job_title=request.job_title,
            company=request.company
        )
        
        # Generate recommendation
        recommendation = get_application_recommendation(match_score)
        
        logger.info(f"Skill gap analysis complete: {match_score}% match")
        
        return SkillGapAnalysisResponse(
            match_score=round(match_score, 1),
            matching_skills=matching_skills,
            missing_skills=missing_skills,
            insights=insights,
            recommendation=recommendation
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in skill gap analysis: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/courses/recommend", response_model=List[CourseRecommendation])
async def recommend_courses(skills: List[str]):
    """
    Get course recommendations for skills.
    
    Args:
        skills: List of skills to find courses for
        
    Returns:
        List of course recommendations
    """
    try:
        recommendations = []
        
        # Course database (in production, use real API)
        course_providers = {
            'Coursera': 'https://www.coursera.org/search?query=',
            'edX': 'https://www.edx.org/search?q=',
            'Udemy': 'https://www.udemy.com/courses/search/?q=',
            'LinkedIn Learning': 'https://www.linkedin.com/learning/search?keywords=',
            'Pluralsight': 'https://www.pluralsight.com/search?q='
        }
        
        for skill in skills[:5]:  # Limit to top 5 skills
            # Generate course recommendations
            for provider, base_url in course_providers.items():
                recommendations.append(
                    CourseRecommendation(
                        provider=provider,
                        title=f"{skill} Complete Course",
                        description=f"Master {skill} with hands-on projects and real-world examples",
                        url=f"{base_url}{skill.replace(' ', '+')}",
                        duration="4-8 weeks",
                        level="Beginner to Advanced"
                    )
                )
        
        return recommendations[:10]  # Return top 10
        
    except Exception as e:
        logger.error(f"Error recommending courses: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health_check():
    """Health check endpoint for container monitoring."""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    }


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def extract_text_from_pdf(file_stream):
    """Extract text from PDF file."""
    try:
        pdf_reader = PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return ""


def extract_text_from_docx(file_stream):
    """Extract text from DOCX file."""
    try:
        doc = Document(file_stream)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return ""


def categorize_skills(skills: List[str]) -> Dict[str, List[str]]:
    """Categorize skills into groups."""
    categories = {
        'Programming Languages': [],
        'Frameworks & Libraries': [],
        'Tools & Platforms': [],
        'Soft Skills': [],
        'Other': []
    }
    
    programming_langs = ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'go', 'rust', 'swift', 'kotlin']
    frameworks = ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'node.js']
    tools = ['git', 'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'jenkins', 'terraform']
    soft = ['leadership', 'communication', 'teamwork', 'problem solving', 'analytical']
    
    for skill in skills:
        skill_lower = skill.lower()
        if any(lang in skill_lower for lang in programming_langs):
            categories['Programming Languages'].append(skill)
        elif any(fw in skill_lower for fw in frameworks):
            categories['Frameworks & Libraries'].append(skill)
        elif any(tool in skill_lower for tool in tools):
            categories['Tools & Platforms'].append(skill)
        elif any(soft_skill in skill_lower for soft_skill in soft):
            categories['Soft Skills'].append(skill)
        else:
            categories['Other'].append(skill)
    
    return {k: v for k, v in categories.items() if v}


def estimate_experience_level(text: str) -> str:
    """Estimate experience level from resume text."""
    text_lower = text.lower()
    
    # Simple heuristic based on keywords
    senior_keywords = ['senior', 'lead', 'principal', 'architect', 'director', 'manager']
    mid_keywords = ['years', 'experience', 'developed', 'led', 'managed']
    
    if any(keyword in text_lower for keyword in senior_keywords):
        return 'Senior'
    elif any(keyword in text_lower for keyword in mid_keywords):
        return 'Mid-Level'
    else:
        return 'Entry-Level'


def generate_insights(
    match_score: float,
    matching_skills: List[str],
    missing_skills: List[str],
    job_title: str,
    company: Optional[str]
) -> List[str]:
    """Generate insights for skill gap analysis."""
    insights = []
    
    if match_score >= 80:
        insights.append(f"You're an excellent match for this {job_title} position!")
        insights.append("Your skills align very well with the job requirements.")
    elif match_score >= 60:
        insights.append(f"You have a good foundation for this {job_title} role.")
        insights.append("Consider highlighting your transferable skills in your application.")
    elif match_score >= 40:
        insights.append("There are some skill gaps to address for this position.")
        insights.append("Focus on building the missing skills through online courses or projects.")
    else:
        insights.append("This role requires significant upskilling.")
        insights.append("Consider applying after gaining more relevant experience.")
    
    if len(matching_skills) > 0:
        top_skills = matching_skills[:3]
        insights.append(f"Highlight these matching skills: {', '.join(top_skills)}")
    
    if len(missing_skills) > 0 and len(missing_skills) <= 3:
        insights.append(f"Focus on learning: {', '.join(missing_skills)}")
    elif len(missing_skills) > 3:
        insights.append(f"Priority skills to learn: {', '.join(missing_skills[:3])}")
    
    return insights


def get_application_recommendation(match_score: float) -> str:
    """Get application recommendation based on match score."""
    if match_score >= 80:
        return "Strongly Recommended - Apply with confidence!"
    elif match_score >= 60:
        return "Recommended - Good match, worth applying"
    elif match_score >= 40:
        return "Consider Carefully - Significant skill gaps exist"
    else:
        return "Not Recommended - Focus on skill development first"


def main():
    """Run the API server."""
    import uvicorn
    
    port = int(os.getenv("API_PORT", 8000))
    host = os.getenv("API_HOST", "0.0.0.0")
    reload = os.getenv("API_RELOAD", "true").lower() == "true"
    
    logger.info(f"Starting API server on {host}:{port}")
    
    uvicorn.run(
        "backend.main:app",
        host=host,
        port=port,
        reload=reload
    )


if __name__ == "__main__":
    main()
